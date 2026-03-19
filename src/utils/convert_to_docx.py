import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        
        # Handle headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
        elif line == '---':
            doc.add_page_break() # Or handle as a separator
        elif line.startswith('- ') or line.startswith('* '):
            # Simple list item
            text = line[2:]
            # Basic bold parsing: **text**
            doc.add_paragraph(text, style='List Bullet')
        elif re.match(r'^\d+\.', line):
            # Numbered list
            text = re.sub(r'^\d+\.\s+', '', line)
            doc.add_paragraph(text, style='List Number')
        elif not line:
            # Skip empty lines but add a small spacing if needed
            continue
        else:
            # Regular paragraph
            # Basic bold parsing: **text**
            # (Very simplistic regex for bold/italic for a utility script)
            clean_text = line.replace('**', '').replace('*', '')
            doc.add_paragraph(clean_text)

    doc.save(docx_path)
    print(f"Successfully saved to {docx_path}")

if __name__ == "__main__":
    md_file = r"C:\Users\IT PC 2\.gemini\antigravity\brain\360f6fdc-d61c-4cf2-8770-4e95d5962261\Telegram_Mini_App_Documentation.md"
    docx_file = r"C:\Users\IT PC 2\.gemini\antigravity\brain\360f6fdc-d61c-4cf2-8770-4e95d5962261\Telegram_Mini_App_Documentation.docx"
    markdown_to_docx(md_file, docx_file)
